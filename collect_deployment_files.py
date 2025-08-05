import os
from docx import Document

# File extensions and specific files to include
INCLUDED_EXTENSIONS = {'.txt', '.yml', '.yaml', '.pyt', '.bat', '.md'}
SPECIFIC_FILES = {
    'requirements.txt',
    'requirements_production.txt', 
    'requirements_dev.txt',
    'requirements_complete.txt',
    'Dockerfile',
    'docker-compose.yml',
    'docker-run.bat',
    'db_manager.pyt',
    'REQUIREMENTS_README.md',
    'DEPLOYMENT.md',
    'DATABASE_FIX_SUMMARY.md'
}

def add_file_to_doc(doc, file_path):
    """Add a file to the Word document with the same format as write_to_word"""
    doc.add_heading(f'File: {file_path}', level=2)
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        doc.add_paragraph(content)
    except Exception as e:
        doc.add_paragraph(f"[Could not read file: {e}]")
    doc.add_paragraph("\n" + "-"*50 + "\n")

def collect_deployment_files_to_doc(root_dir, output_path):
    """Collect deployment-related files into a Word document"""
    doc = Document()
    doc.add_heading('ArcSpatialDB Deployment Files Collection', 0)
    doc.add_paragraph('This document contains all deployment-related files including requirements, Docker configuration, and the ArcGIS Pro plugin.')
    doc.add_paragraph("\n" + "="*80 + "\n")

    # Track found files
    found_files = []
    missing_files = []

    # First, collect specific files we want
    for filename in SPECIFIC_FILES:
        file_path = os.path.join(root_dir, filename)
        if os.path.exists(file_path):
            print(f"📄 Adding: {filename}")
            add_file_to_doc(doc, file_path)
            found_files.append(filename)
        else:
            print(f"⚠️ Missing: {filename}")
            missing_files.append(filename)

    # Then, look for any other files with included extensions
    for dirpath, _, filenames in os.walk(root_dir):
        for filename in filenames:
            _, ext = os.path.splitext(filename)
            if ext.lower() in INCLUDED_EXTENSIONS:
                full_path = os.path.join(dirpath, filename)
                rel_path = os.path.relpath(full_path, root_dir)
                
                # Skip if we already processed this file
                if filename in found_files:
                    continue
                
                # Skip if it's in a subdirectory we don't want
                if any(skip_dir in rel_path for skip_dir in ['.git', '__pycache__', '.idea', '.pytest_cache']):
                    continue
                
                print(f"📄 Adding additional: {rel_path}")
                add_file_to_doc(doc, full_path)
                found_files.append(filename)

    # Add summary at the end
    doc.add_heading('Collection Summary', level=1)
    doc.add_paragraph(f"Total files collected: {len(found_files)}")
    
    if found_files:
        doc.add_paragraph("Files included:")
        for filename in sorted(found_files):
            doc.add_paragraph(f"• {filename}", style='List Bullet')
    
    if missing_files:
        doc.add_paragraph("Files not found:")
        for filename in sorted(missing_files):
            doc.add_paragraph(f"• {filename}", style='List Bullet')

    # Save the document
    doc.save(output_path)
    print(f"\n✅ Document saved to: {output_path}")
    print(f"📊 Total files collected: {len(found_files)}")
    
    if missing_files:
        print(f"⚠️ Missing files: {len(missing_files)}")
        for filename in missing_files:
            print(f"   - {filename}")

# Path configuration
root_directory = r"C:\Users\yuval\PycharmProjects\ArcSpatialDB"
output_docx = r"C:\Users\yuval\PycharmProjects\ArcSpatialDB\deployment_files_collected.docx"

if __name__ == "__main__":
    print("🚀 Collecting ArcSpatialDB Deployment Files")
    print("=" * 50)
    collect_deployment_files_to_doc(root_directory, output_docx)
    print("\n🎉 Collection completed!") 